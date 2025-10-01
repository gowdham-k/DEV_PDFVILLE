from pdf2image import convert_from_bytes 
import pytesseract 
from docx import Document 
from flask import request, jsonify, send_file 
import io, os, tempfile, traceback 
import shutil  # ✅ added for auto-detect
from restrictions import check_convert_pdf_restrictions
from utils import create_temp_file, cleanup_file

# ✅ Auto-detect tesseract.exe from PATH
tesseract_path = shutil.which("tesseract")
if tesseract_path:
    pytesseract.pytesseract.tesseract_cmd = tesseract_path
else:
    raise FileNotFoundError(
        "Tesseract not found. Please install it and add to PATH: https://github.com/UB-Mannheim/tesseract/wiki"
    )

def handle_convert_pdf_to_word(): 
    try: 
        uploaded_file = request.files.get("file") 
        if not uploaded_file: 
            return jsonify({"error": "No file provided"}), 400 
 
        # Save uploaded file temporarily for restriction check
        temp_input_path = create_temp_file(uploaded_file.filename)
        uploaded_file.save(temp_input_path)
        
        # Check restrictions for PDF files
        email = request.form.get("email", "anonymous@example.com")
        restriction_error = check_convert_pdf_restrictions(email, [temp_input_path])
        if restriction_error:
            cleanup_file(temp_input_path)
            return jsonify(restriction_error), 403

        pdf_bytes = uploaded_file.read() 
        images = convert_from_bytes(pdf_bytes)  # convert pdf → images 
 
        # Clean up temp input file
        cleanup_file(temp_input_path) 
 
        doc = Document() 
 
        for i, image in enumerate(images, start=1): 
            text = pytesseract.image_to_string(image) 
            doc.add_paragraph(f"--- Page {i} ---") 
            doc.add_paragraph(text) 
            doc.add_page_break() 
 
        temp_dir = tempfile.mkdtemp() 
        output_path = os.path.join(temp_dir, "output.docx") 
        doc.save(output_path) 
 
        return send_file(output_path, as_attachment=True, download_name="converted.docx") 
 
    except Exception as e: 
        print("🔥 ERROR in /convert-word route:") 
        traceback.print_exc() 
        return jsonify({"error": str(e)}), 500
